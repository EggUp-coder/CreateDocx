import os
import re
import shutil
from docx import Document
from openpyxl import load_workbook
from tkinter import filedialog
import easygui
from win32com.shell import shell, shellcon

def get_path():
    door = 0
    path_root = filedialog.askdirectory()
    list_files = os.listdir(path_root)
    num_doc = 0
    num_excel = 0
    for i in list_files:
        if i[-5:] == '.docx' :
            path_doc = path_root + '//' + i
            num_doc += 1
        elif i[-5:] == '.xlsx': 
            path_excel = path_root + '//' + i
            num_excel += 1
    
    if num_doc * num_excel == 0:
        easygui.msgbox('缺少docx格式的word或xlsx格式的excel。')
    elif num_doc * num_excel > 1: 
        easygui.msgbox('有其他的word或excel，请移除。')
    elif num_doc * num_excel == 1:
        door = 1

    return door,path_doc,path_excel


def info_update(path_doc, old_info, new_info):
    doc = Document(path_doc)
    for para in doc.paragraphs:
        for run in para.runs:
            run.text = run.text.replace(old_info, new_info)
            run.text = re.sub(r'[《》]', '', run.text)
    
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                cell.text = cell.text.replace(old_info, new_info)
    doc.save(path_doc)
  
door,path_doc,path_excel = get_path()

if door == 1:
    ilist =shell.SHGetSpecialFolderLocation(0, shellcon.CSIDL_DESKTOP)
    path_desktop = shell.SHGetPathFromIDList(ilist).decode()
    path_new_folders = path_desktop+'\\'+'批量生成'
        
    while os.path.exists(path_new_folders) :
        path_new_folders +='-新'
    
    os.makedirs(path_new_folders)

    list_new_doc_path = []

    wb = load_workbook(path_excel)
    ws = wb.get_sheet_by_name('Sheet1')
    
    for i in range(2,1000):
        if ws.cell(i,1).value == None:
            max_row = i-1
            break
        else :
            new_doc_path = path_new_folders+"//"+ws.cell(i,2).value + '-' +ws.cell(i,3).value+'.docx'
            shutil.copy(path_doc,new_doc_path)
            list_new_doc_path.append(new_doc_path)

    for i in range(2,100):
        if ws.cell(1,i).value == None:
            max_col = i-1 
            break 

    for row in range(len(list_new_doc_path)):
        for col in range(2,max_col+1):
            path_doc_new = list_new_doc_path[row]

            old_info = str(ws.cell(1,col).value)

            new_info = str(ws.cell(row+2,col).value)

            info_update(path_doc_new, old_info, new_info)